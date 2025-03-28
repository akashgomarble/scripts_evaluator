import os
import pandas as pd
import PyPDF2
from docx import Document
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.tag import pos_tag
from nltk.corpus import stopwords
from nltk.chunk import ne_chunk
from nltk.tag import pos_tag_sents
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np
from typing import List, Dict, Union, Any
import json
from pydantic import BaseModel, Field
from openai import OpenAI
import ssl
import logging
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize OpenAI client with API key from environment

# Get model name from en
api_key = os.getenv("OPENAI_API_KEY")
if not api_key:
    raise ValueError("OPENAI_API_KEY environment variable is not set")

# Get model name from environment
model_name = os.getenv("MODEL_NAME", "gpt-4")  # Default to gpt-4 if not set
openai_client = OpenAI(api_key = api_key)

# SSL certificate verification workaround for NLTK
try:
    _create_unverified_https_context = ssl._create_unverified_context
except AttributeError:
    pass
else:
    ssl._create_default_https_context = _create_unverified_https_context

class ScriptEvaluator:
    def __init__(self):
        # Download required NLTK data with SSL verification disabled
        try:
            nltk.download('punkt', quiet=True)
            nltk.download('stopwords', quiet=True)
            nltk.download('averaged_perceptron_tagger', quiet=True)
            nltk.download('maxent_ne_chunker', quiet=True)
            nltk.download('words', quiet=True)
        except Exception as e:
            logger.warning(f"Error downloading NLTK data: {str(e)}")
            logger.warning("Some NLTK features may not be available")
        
        # Use the global OpenAI client instance
        self.openai_client = openai_client
        self.model_name = model_name
        
        # Define evaluation criteria schema
        class EvaluationScore(BaseModel):
            score: float = Field(description="Score between 0 and 1")
            reasoning: str = Field(description="Explanation for the score")
        
        self.evaluation_schema = EvaluationScore
        
        # Define evaluation instructions
        self.evaluation_instructions = {
            "content_quality": """Evaluate the content quality of the script based on:
            - Vocabulary and grammar
            - Writing style
            - Language proficiency
            Provide a score between 0 and 1 with reasoning.""",
            
            "structure": """Evaluate the structure and organization of the script based on:
            - Document structure
            - Section organization
            - Flow and coherence
            Provide a score between 0 and 1 with reasoning.""",
            
            "technical_accuracy": """Evaluate the technical accuracy of the script based on:
            - Factual correctness
            - Reference alignment
            - Technical precision
            Provide a score between 0 and 1 with reasoning.""",
            
            "clarity": """Evaluate the clarity and readability of the script based on:
            - Clear expression
            - Readability
            - Audience appropriateness
            Provide a score between 0 and 1 with reasoning.""",
            
            "completeness": """Evaluate the completeness of the script based on:
            - Required sections
            - Content coverage
            - Missing elements
            Provide a score between 0 and 1 with reasoning.""",
            
            "originality": """Evaluate the originality of the script based on:
            - Unique content
            - Creative elements
            - Reference differentiation
            Provide a score between 0 and 1 with reasoning."""
        }
        
    def parse_document(self, file_path: str) -> str:
        """Parse a document file and return its text content."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Get file extension
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        
        try:
            if ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            
            elif ext == '.csv':
                df = pd.read_csv(file_path)
                return df.to_string()
            
            elif ext == '.docx':
                doc = Document(file_path)
                return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            
            elif ext == '.pdf':
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    return ' '.join([page.extract_text() for page in reader.pages])
            
            else:
                raise ValueError(f"Unsupported file format: {ext}")
                
        except Exception as e:
            logger.error(f"Error parsing document {file_path}: {str(e)}")
            raise

    def evaluate_script(self, script_path: str, reference_docs: List[str] = None) -> Dict[str, Any]:
        """Comprehensive script evaluation using LangSmith."""
        # Parse script
        script_text = self.parse_document(script_path)
        
        # Convert script to LangChain Document format
        script_doc = LangchainDocument(
            page_content=script_text,
            metadata={"source": script_path}
        )
        
        # Split text into chunks for evaluation
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200
        )
        script_chunks = text_splitter.split_documents([script_doc])
        
        # Create dataset for evaluation
        dataset = self.client.create_dataset(
            dataset_name="Script Evaluation Dataset",
            description="Dataset for script evaluation"
        )
        
        # Prepare inputs and outputs for evaluation
        inputs = [{"text": chunk.page_content} for chunk in script_chunks]
        outputs = []
        
        # Run evaluations for each category
        evaluation_results = {}
        for category, instructions in self.evaluation_instructions.items():
            result = self._evaluate_category(script_text, category, instructions)
            evaluation_results[category] = result
        
        return evaluation_results

    def _evaluate_category(self, text: str, category: str, instructions: str) -> Dict[str, Any]:
        """Evaluate a specific category using OpenAI."""
        try:
            logger.info(f"Evaluating {category} using model {self.model_name}")
            response = self.openai_client.chat.completions.create(
                model=self.model_name,
                messages=[
                    {"role": "system", "content": "You are an expert at extracting structured information from text."},
                    {"role": "user", "content": f"Script text: {text}"}
                ],
                temperature=0
            )
            
            content = response.choices[0].message.content
            logger.debug(f"Raw response for {category}: {content}")
            
            # Parse the response
            try:
                if "Score:" in content:
                    parts = content.split("Score:", 1)
                    score_text = parts[1].split("\n")[0].strip()
                    score = float(score_text)
                    reasoning = parts[1].split("\n", 1)[1].strip() if len(parts[1].split("\n")) > 1 else parts[0].strip()
                else:
                    # If no explicit score, try to find a number in the first line
                    first_line = content.split("\n")[0]
                    import re
                    score_match = re.search(r'\b(\d+(\.\d+)?)\b', first_line)
                    score = float(score_match.group(1)) if score_match else 0.7
                    reasoning = content.strip()
                
                logger.info(f"Evaluation complete for {category}. Score: {score}")
                return {
                    'metrics': {
                        'score': score,
                        'reasoning': reasoning
                    },
                    'score': score,
                    'feedback': [reasoning]
                }
            except Exception as e:
                logger.warning(f"Error parsing response for {category}: {str(e)}")
                return {
                    'metrics': {
                        'score': 0.7,
                        'reasoning': content
                    },
                    'score': 0.7,
                    'feedback': [content]
                }
                
        except Exception as e:
            logger.error(f"OpenAI API error for {category}: {str(e)}")
            return {
                'metrics': {
                    'score': 0,
                    'reasoning': str(e)
                },
                'score': 0,
                'feedback': [f"Error: {str(e)}"]
            }

    def generate_report(self, evaluation_results: Dict[str, Any], output_path: str = None):
        """Generate a detailed evaluation report."""
        # Calculate overall score
        overall_score = sum(
            category['score'] for category in evaluation_results.values()
        ) / len(evaluation_results)
        
        # Collect all feedback
        all_feedback = []
        for category, results in evaluation_results.items():
            all_feedback.extend(results['feedback'])
        
        report = {
            'overall_score': overall_score,
            'category_scores': {
                category: results['score']
                for category, results in evaluation_results.items()
            },
            'detailed_metrics': {
                category: results['metrics']
                for category, results in evaluation_results.items()
            },
            'feedback': all_feedback,
            'recommendations': self._generate_recommendations(evaluation_results)
        }
        
        if output_path:
            with open(output_path, 'w') as f:
                json.dump(report, f, indent=4)
        
        return report

    def _generate_recommendations(self, results: Dict[str, Any]) -> List[str]:
        """Generate overall recommendations based on all evaluation results."""
        recommendations = []
        
        # Add recommendations based on lowest scoring categories
        category_scores = {
            category: results[category]['score']
            for category in results
        }
        
        lowest_categories = sorted(category_scores.items(), key=lambda x: x[1])[:2]
        for category, score in lowest_categories:
            if score < 0.6:
                recommendations.append(f"Focus on improving {category.replace('_', ' ')}")
        
        return recommendations 